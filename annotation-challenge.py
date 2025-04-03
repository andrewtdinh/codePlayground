import re
import requests
import pandas as pd
from html.parser import HTMLParser
from urllib.parse import urlparse, parse_qs

class TableParser(HTMLParser):
    """
    Custom HTML parser to extract table data from Google Docs HTML export
    """
    def __init__(self):
        super().__init__()
        self.tables = []
        self.current_table = []
        self.current_row = []
        self.current_cell = ""
        self.in_table = False
        self.in_row = False
        self.in_cell = False
        
    def handle_starttag(self, tag, attrs):
        if tag == 'table':
            self.in_table = True
            self.current_table = []
        elif tag == 'tr' and self.in_table:
            self.in_row = True
            self.current_row = []
        elif (tag == 'td' or tag == 'th') and self.in_row:
            self.in_cell = True
            self.current_cell = ""
            
    def handle_endtag(self, tag):
        if tag == 'table':
            self.tables.append(self.current_table)
            self.in_table = False
        elif tag == 'tr' and self.in_table:
            self.current_table.append(self.current_row)
            self.in_row = False
        elif (tag == 'td' or tag == 'th') and self.in_row:
            self.current_row.append(self.current_cell.strip())
            self.in_cell = False
            
    def handle_data(self, data):
        if self.in_cell:
            self.current_cell += data

def print_unicode_grid_from_gdoc(url):
    """
    Retrieves data from a Google Doc table and prints a grid of Unicode characters.
    
    Args:
        url (str): URL of the Google Doc containing a table with columns:
                  'x-coordinate', 'Character', and 'y-coordinate'
    """
    try:
        # Extract document ID from URL
        doc_id = extract_doc_id(url)
        
        # Try multiple formats to extract data
        coordinates = []
        
        # Try HTML format first (more structured)
        coordinates = try_html_format(doc_id)
        
        # If HTML didn't work, try plain text
        if not coordinates:
            coordinates = try_text_format(doc_id)
            
        # If still no coordinates, try downloading the document as DOCX
        if not coordinates:
            coordinates = try_docx_format(doc_id)
        
        if not coordinates:
            print("No valid coordinate data found in the document.")
            return
        
        # Convert to DataFrame for easier processing
        df = pd.DataFrame(coordinates, columns=['x-coordinate', 'Character', 'y-coordinate'])
        
        # Ensure coordinate columns are numeric and strip any non-numeric characters
        df['x-coordinate'] = pd.to_numeric(df['x-coordinate'].str.replace(r'[^\d]', '', regex=True))
        df['y-coordinate'] = pd.to_numeric(df['y-coordinate'].str.replace(r'[^\d]', '', regex=True))
        
        # Calculate grid dimensions
        max_x = int(df['x-coordinate'].max())
        max_y = int(df['y-coordinate'].max())
        
        # Create an empty grid filled with spaces
        grid = [[' ' for _ in range(max_x + 1)] for _ in range(max_y + 1)]
        
        # Place characters in the grid
        for _, row in df.iterrows():
            x = int(row['x-coordinate'])
            y = int(row['y-coordinate'])
            char = row['Character']
            if len(char) > 0:  # Ensure we have a character
                grid[y][x] = char[0]  # Take first character if multiple
        
        # Print the grid
        for row in grid:
            print(''.join(row))
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()


def extract_doc_id(url):
    """Extract the document ID from a Google Docs URL."""
    # Parse URL patterns like /document/d/{doc_id}/edit or /document/u/0/d/{doc_id}/edit
    patterns = [
        r'/document/d/([a-zA-Z0-9-_]+)',
        r'/document/u/\d+/d/([a-zA-Z0-9-_]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    
    # Try to extract from query parameters (like ?id={doc_id})
    parsed_url = urlparse(url)
    query_params = parse_qs(parsed_url.query)
    if 'id' in query_params:
        return query_params['id'][0]
    
    raise ValueError("Could not extract document ID from URL")


def try_html_format(doc_id):
    """Try to extract coordinates from HTML export."""
    coordinates = []
    
    try:
        # Use HTML export
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"
        response = requests.get(export_url)
        
        if response.status_code == 200:
            html_content = response.text
            
            # Parse HTML to find tables
            parser = TableParser()
            parser.feed(html_content)
            
            # Process each table found
            for table in parser.tables:
                if len(table) < 2:  # Need at least header + one data row
                    continue
                
                # Check if header row contains our expected columns
                header = [cell.lower() for cell in table[0]]
                
                # Find indices for our columns (allowing for variations in naming)
                x_idx = -1
                char_idx = -1
                y_idx = -1
                
                for i, cell in enumerate(header):
                    if 'x' in cell and ('coord' in cell or 'position' in cell):
                        x_idx = i
                    elif 'char' in cell:
                        char_idx = i
                    elif 'y' in cell and ('coord' in cell or 'position' in cell):
                        y_idx = i
                
                # If we found all three columns
                if x_idx >= 0 and char_idx >= 0 and y_idx >= 0:
                    # Process data rows
                    for row in table[1:]:
                        if len(row) > max(x_idx, char_idx, y_idx):
                            x = row[x_idx].strip()
                            char = row[char_idx].strip()
                            y = row[y_idx].strip()
                            
                            # Check if values look valid
                            if re.search(r'\d+', x) and re.search(r'\d+', y) and char:
                                coordinates.append([x, char, y])
                    
                    # If we found valid data, return it
                    if coordinates:
                        return coordinates
    
    except Exception as e:
        print(f"HTML extraction error: {e}")
    
    return coordinates


def try_text_format(doc_id):
    """Try to extract coordinates from plain text export."""
    coordinates = []
    
    try:
        # Use plain text export
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        response = requests.get(export_url)
        
        if response.status_code == 200:
            text_content = response.text
            
            # Try to identify table structure
            lines = text_content.split('\n')
            
            # Look for lines with "x-coordinate" and "y-coordinate"
            table_start = -1
            for i, line in enumerate(lines):
                if ('x-coordinate' in line.lower() and 
                    'character' in line.lower() and 
                    'y-coordinate' in line.lower()):
                    table_start = i
                    break
            
            if table_start >= 0:
                # Extract data using pattern matching
                for i in range(table_start + 1, len(lines)):
                    line = lines[i].strip()
                    if not line:
                        continue
                    
                    # Try to extract three columns of data
                    # Pattern: numbers followed by character followed by numbers
                    # More flexible version: Look for numbers and single character
                    match = re.search(r'(\d+)\s+(\S)[\s\D]*(\d+)', line)
                    if match:
                        x = match.group(1)
                        char = match.group(2)
                        y = match.group(3)
                        coordinates.append([x, char, y])
    
    except Exception as e:
        print(f"Text extraction error: {e}")
    
    return coordinates


def try_docx_format(doc_id):
    """
    Try to extract coordinates from DOCX format.
    Requires the python-docx library.
    """
    coordinates = []
    
    try:
        # Try to import docx library
        import docx
        
        # Download the document as DOCX
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
        response = requests.get(export_url)
        
        if response.status_code == 200:
            # Save to temporary file
            with open('temp_doc.docx', 'wb') as f:
                f.write(response.content)
            
            # Open the document
            doc = docx.Document('temp_doc.docx')
            
            # Extract tables
            for table in doc.tables:
                if len(table.rows) < 2:  # Need at least header + one data row
                    continue
                
                # Get header row
                header = [cell.text.lower().strip() for cell in table.rows[0].cells]
                
                # Find our columns
                x_idx = -1
                char_idx = -1
                y_idx = -1
                
                for i, cell in enumerate(header):
                    if 'x' in cell and ('coord' in cell or 'position' in cell):
                        x_idx = i
                    elif 'char' in cell:
                        char_idx = i
                    elif 'y' in cell and ('coord' in cell or 'position' in cell):
                        y_idx = i
                
                # If we found all three columns
                if x_idx >= 0 and char_idx >= 0 and y_idx >= 0:
                    # Process data rows
                    for row in table.rows[1:]:
                        cells = row.cells
                        if len(cells) > max(x_idx, char_idx, y_idx):
                            x = cells[x_idx].text.strip()
                            char = cells[char_idx].text.strip()
                            y = cells[y_idx].text.strip()
                            
                            # Check if values look valid
                            if re.search(r'\d+', x) and re.search(r'\d+', y) and char:
                                coordinates.append([x, char, y])
            
            # Clean up
            import os
            if os.path.exists('temp_doc.docx'):
                os.remove('temp_doc.docx')
    
    except ImportError:
        print("python-docx library not available, skipping DOCX extraction")
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        
    return coordinates


def try_direct_csv_export(doc_id):
    """
    As a fallback, try to access the document as a CSV
    (this works if the document is actually a Google Sheet)
    """
    coordinates = []
    
    try:
        # Try to export as CSV (for Google Sheets)
        export_url = f"https://docs.google.com/spreadsheets/d/{doc_id}/export?format=csv"
        response = requests.get(export_url)
        
        if response.status_code == 200:
            # Parse CSV data
            import csv
            from io import StringIO
            
            csv_data = response.text
            csv_reader = csv.reader(StringIO(csv_data))
            
            # Get header row
            header = next(csv_reader, [])
            header = [col.lower().strip() for col in header]
            
            # Find our columns
            x_idx = -1
            char_idx = -1
            y_idx = -1
            
            for i, col in enumerate(header):
                if 'x' in col and ('coord' in col or 'position' in col):
                    x_idx = i
                elif 'char' in col:
                    char_idx = i
                elif 'y' in col and ('coord' in col or 'position' in col):
                    y_idx = i
            
            # If we found all three columns
            if x_idx >= 0 and char_idx >= 0 and y_idx >= 0:
                # Process data rows
                for row in csv_reader:
                    if len(row) > max(x_idx, char_idx, y_idx):
                        x = row[x_idx].strip()
                        char = row[char_idx].strip()
                        y = row[y_idx].strip()
                        
                        # Check if values look valid
                        if re.search(r'\d+', x) and re.search(r'\d+', y) and char:
                            coordinates.append([x, char, y])
    
    except Exception as e:
        print(f"CSV extraction error: {e}")
        
    return coordinates


def manual_grid_extraction(doc_id):
    """
    Last resort: Try to manually extract a grid from the document text
    by finding patterns that look like coordinates and characters.
    """
    coordinates = []
    
    try:
        # Use plain text export
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        response = requests.get(export_url)
        
        if response.status_code == 200:
            text_content = response.text
            
            # Try to find all patterns that look like coordinates and characters
            # Pattern: (x,y)=char or similar patterns
            patterns = [
                r'.*?(\d+).*?(\d+).*?[=:].*?(\S)',  # Patterns like "x=1, y=2: *"
                r'\((\d+),\s*(\d+)\)\s*[=:]?\s*(\S)',  # Patterns like "(1, 2) = *"
                r'(\d+)[,\s]+(\S)[,\s]+(\d+)',  # Patterns like "1, *, 2"
            ]
            
            for pattern in patterns:
                matches = re.findall(pattern, text_content)
                if matches:
                    for match in matches:
                        if len(match) == 3:
                            # Determine which values are x, y, and char
                            if len(match[1]) == 1 and not match[1].isdigit():
                                # Format is likely (x, char, y)
                                x, char, y = match
                            else:
                                # Format is likely (x, y, char)
                                x, y, char = match
                            
                            coordinates.append([x, char, y])
    
    except Exception as e:
        print(f"Manual extraction error: {e}")
        
    return coordinates


if __name__ == "__main__":
    # Test with a sample URL
    url = "https://docs.google.com/document/d/1yiZrWz0-1bqhvk4Rsrcj1TU5SjPn3_KKoYRNA40gJa4/edit"
    print_unicode_grid_from_gdoc(url)
